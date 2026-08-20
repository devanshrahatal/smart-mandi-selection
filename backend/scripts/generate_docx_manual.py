import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn
import os

def create_manual():
    doc = docx.Document()

    # Set standard margins (1 inch)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.85)
        section.right_margin = Inches(0.85)

    def set_cell_background(cell, fill_hex):
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
        cell._tc.get_or_add_tcPr().append(shading)

    def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
        tcPr = cell._tc.get_or_add_tcPr()
        tcMar = OxmlElement('w:tcMar')
        for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
            node = OxmlElement(f'w:{m}')
            node.set(qn('w:w'), str(val))
            node.set(qn('w:type'), 'dxa')
            tcMar.append(node)
        tcPr.append(tcMar)

    # Palette
    COLOR_PRIMARY = RGBColor(16, 185, 129)   # Emerald 500
    COLOR_DARK = RGBColor(15, 23, 42)        # Slate 900
    COLOR_MUTED = RGBColor(100, 116, 139)    # Slate 500
    COLOR_TEXT = RGBColor(30, 41, 59)        # Slate 800

    # Document Header Title
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run("🌾 SMART MANDI SELECTION PLATFORM")
    title_run.font.name = 'Calibri'
    title_run.font.size = Pt(22)
    title_run.font.bold = True
    title_run.font.color.rgb = COLOR_PRIMARY

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_p.add_run("Comprehensive End-to-End User Manual & Step-by-Step Operating Guide\n")
    sub_run.font.name = 'Calibri'
    sub_run.font.size = Pt(13)
    sub_run.font.italic = True
    sub_run.font.color.rgb = COLOR_MUTED

    # Meta Table
    meta_table = doc.add_table(rows=4, cols=2)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_data = [
        ("Project Lead & Author:", "Devansh Rahatal (Lead Architect & Developer)"),
        ("Project Classification:", "Agricultural Price Intelligence & Logistics Optimization"),
        ("Web Application:", "React (Vite) + FastAPI + Leaflet GIS + Twilio WhatsApp Bot"),
        ("Documentation Version:", "Version 1.0 (Production & Jury Presentation Edition)"),
    ]
    for i, (label, val) in enumerate(meta_data):
        row = meta_table.rows[i]
        c1, c2 = row.cells[0], row.cells[1]
        c1.width = Inches(2.2)
        c2.width = Inches(4.5)
        set_cell_background(c1, "F1F5F9")
        set_cell_background(c2, "F8FAFC")
        set_cell_margins(c1, 80, 80, 120, 120)
        set_cell_margins(c2, 80, 80, 120, 120)

        p1 = c1.paragraphs[0]
        r1 = p1.add_run(label)
        r1.font.bold = True
        r1.font.size = Pt(9.5)
        r1.font.color.rgb = COLOR_DARK

        p2 = c2.paragraphs[0]
        r2 = p2.add_run(val)
        r2.font.size = Pt(9.5)
        r2.font.color.rgb = COLOR_TEXT

    doc.add_paragraph()

    # Section Helper
    def add_heading_1(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(text)
        r.font.name = 'Calibri'
        r.font.size = Pt(15)
        r.font.bold = True
        r.font.color.rgb = COLOR_DARK
        return p

    def add_heading_2(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(text)
        r.font.name = 'Calibri'
        r.font.size = Pt(12)
        r.font.bold = True
        r.font.color.rgb = COLOR_PRIMARY
        return p

    def add_p(text, bold_prefix=None):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        if bold_prefix:
            r_pre = p.add_run(bold_prefix)
            r_pre.font.bold = True
            r_pre.font.size = Pt(10)
            r_pre.font.color.rgb = COLOR_DARK
        r = p.add_run(text)
        r.font.size = Pt(10)
        r.font.color.rgb = COLOR_TEXT
        return p

    # --- 1. EXECUTIVE SUMMARY ---
    add_heading_1("1. Executive Summary & The Core Problem")
    add_p("Most agricultural market applications display only the Gross Modal Price (advertised mandi price). As a result, farmers frequently travel long distances chasing high prices, only to suffer severe financial losses due to transport fees, APMC mandi commissions, handling charges, and transit spoilage.")
    add_p("Smart Mandi solves this 'Price Illusion' by computing the exact Net Take-Home Profit for every candidate mandi within a 300 km radius.")
    
    add_p("Net Profit = Modal Price - Transport Cost - Loading/Unloading - Mandi Commission - Transit Spoilage", bold_prefix="Mathematical Formula: ")
    add_p("Where Spoilage Deduction = Modal Price × Perishability Index (Tomato: 0.85, Onion: 0.25, Potato: 0.20, Wheat: 0.05, Banana: 0.70) × (Transit Hours / 24) × 0.15.")

    # --- 2. WEB APPLICATION WALKTHROUGH ---
    add_heading_1("2. Web Application Step-by-Step User Guide")

    add_heading_2("2.1 Public Landing Page (Home)")
    add_p("1. Navigate to the web application URL (e.g. http://localhost:5173 or the deployed Vercel URL).")
    add_p("2. View the Live Scenario Comparison Table illustrating how a Jaipur tomato farmer makes ₹169/quintal MORE net cash in nearby Kota than in Azadpur (Delhi), despite Azadpur having a higher raw price.")
    add_p("3. Switch language effortlessly between English, Hindi (हिंदी), Marathi (मराठी), and Gujarati (ગુજરાતી) using the top language selector.")
    add_p("4. Access the 1-Click WhatsApp Assistant banner or launch the In-Browser Voice Simulator.")

    add_heading_2("2.2 Interactive Profit Map (/map)")
    add_p("1. Click 'Profit Map' in the navigation bar.")
    add_p("2. Select your Origin Hub (e.g., Vadodara, Surat, Rajkot, Jaipur, Pune, Nashik, etc.).")
    add_p("3. Select your Crop (Tomato, Onion, Potato, Wheat, Banana) and set your Harvest Quantity via the interactive slider.")
    add_p("4. Observe the Leaflet GIS Map: It automatically plots ranked profit routes from your farm:")
    add_p("   • Green Route (#1): Maximum Net Profit destination (Optimal market).")
    add_p("   • Teal Route (#2): Moderate profit alternative.")
    add_p("   • Amber/Red Route (#3): Lower net profit / high-distance market.")
    add_p("5. Click any Mandi Pin to view a full breakdown modal (Take-Home Net Profit, Gross Price, Total Deductions, Distance in km).")

    add_heading_2("2.3 Kisan Pool Shared Freight Optimizer (/pooling)")
    add_p("1. Click 'Kisan Pool' in the navigation bar.")
    add_p("2. Use the Interactive Savings Calculator:")
    add_p("   • Set Solo Quantity (e.g., 12 quintals).")
    add_p("   • Set Total Pooled Quantity (e.g., 36 quintals).")
    add_p("   • Set Distance (e.g., 248 km).")
    add_p("3. The system dynamically matches the optimal commercial vehicle (Tata Ace vs Eicher 14ft vs 19ft Heavy Truck vs 16-Ton Taurus) and calculates instant savings per quintal (typically 35% to 55% reduction in freight costs).")
    add_p("4. View Active Pooling Batches in real time and connect with cluster FPOs with 1 click.")

    add_heading_2("2.4 Admin Intelligence Dashboard (/admin/dashboard)")
    add_p("1. Click 'Admin Dashboard' or navigate to /admin/login.")
    add_p("2. Sign in with the seeded credentials: Username: admin | Password: admin123.")
    add_p("3. View real-time platform metrics: Mandis Tracked, Active Farmers, Average Savings per Quintal, and Total Queries.")
    add_p("4. Inspect the Top Queried Crops volume chart and Top Recommended Mandis.")
    add_p("5. Monitor the Live Farmer Query Stream containing real-time WhatsApp incoming queries and recommendation responses.")
    add_p("6. Click 'Export Report' to download the complete platform dataset as a CSV spreadsheet.")

    add_heading_2("2.5 Mandi Details & Cost Parameters (/admin/mandis, /admin/costs)")
    add_p("1. In 'Price Trends', view 30-day historical modal price curves across crops and markets.")
    add_p("2. In 'Cost Parameters', adjust mandi-specific commission % (e.g. 5% to 8%), loading charges (₹/quintal), unloading charges, and base transport rates (₹/km/q).")
    add_p("3. Changes immediately update the live net profit calculations across the platform.")

    # --- 3. WHATSAPP BOT USER GUIDE ---
    add_heading_1("3. WhatsApp Voice & Text AI Assistant (Step-by-Step)")
    add_p("The WhatsApp Bot allows smallholder farmers to get instant mandi recommendations without downloading or installing any mobile application.")

    add_heading_2("3.1 Connecting to the WhatsApp Gateway")
    add_p("1. Save or open the Twilio WhatsApp Gateway number on your phone: +1 (415) 523-8886.")
    add_p("2. Send the one-time Sandbox Activation Code: join unusual-sea.")
    add_p("3. You will receive an immediate confirmation message: 'You are all set! You can now test your integration.'")

    add_heading_2("3.2 Sending Natural Text Queries")
    add_p("Simply type any natural query. For example:")
    add_p("• 'Tomato 20q from Jaipur'")
    add_p("• 'कांदा 15 क्विंटल नाशिक'")
    add_p("• 'Potato 30q from Agra'")
    add_p("• 'Wheat 50 quintal'")
    add_p("The bot calculates all candidate mandis and returns the ranked net profit table directly in WhatsApp.")

    add_heading_2("3.3 Sending Regional Voice Notes (Speech-to-Speech)")
    add_p("1. Hold the microphone button in WhatsApp and speak in your native dialect (Hindi, Marathi, Gujarati, or English):")
    add_p("   • Hindi: 'भैया 20 क्विंटल टमाटर बेचना है जयपुर से'")
    add_p("   • Marathi: 'नमस्कार, 15 क्विंटल कांदा नाशिकमधून विकायचा आहे'")
    add_p("   • Gujarati: 'નમસ્તે, 20 ક્વિન્ટલ ડુંગળી રાજકોટથી વેચવી છે'")
    add_p("2. The backend uses speech recognition to parse the crop and quantity, and sends back both a text summary AND an instant Voice Note MP3 in your native language explaining the best mandi to visit.")

    add_heading_2("3.4 Sharing GPS Location Pins")
    add_p("1. Tap the attachment icon in WhatsApp -> Select 'Location' -> Send your Current Location 📍.")
    add_p("2. The bot instantly resolves your exact latitude & longitude and computes road distances to all nearby APMC mandis via Google Maps / Haversine routing.")

    # --- 4. TECHNICAL ARCHITECTURE & DEPLOYMENT ---
    add_heading_1("4. Technical Setup & Deployment Reference")
    add_p("• Backend Framework: FastAPI (Python 3.11+), SQLAlchemy 2.0 ORM, Uvicorn.")
    add_p("• Frontend Framework: React 18, Vite 5/6, TailwindCSS, Leaflet Map GIS, Lucide Icons.")
    add_p("• Database: SQLite (Zero-configuration local / Render cloud fallback) or MySQL 8.0.")
    add_p("• Voice AI Engine: gTTS Regional Audio Synthesis + Browser SpeechRecognition.")
    add_p("• WhatsApp Integration: Twilio Programmable Messaging Webhook API.")

    add_heading_2("4.1 Running Locally on Developer Machines")
    add_p("1. Backend: Open terminal in backend/ -> Run: .\\.venv\\Scripts\\python.exe -m uvicorn app.main:app --reload --port 8000")
    add_p("2. Frontend: Open terminal in frontend/ -> Run: npm run dev -> Open http://localhost:5173")
    add_p("3. Interactive Swagger API Docs: Available at http://localhost:8000/docs")

    # Output file
    output_path = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", "Smart_Mandi_Platform_User_Manual.docx"))
    doc.save(output_path)
    print(f"Successfully generated User Manual at: {output_path}")

if __name__ == "__main__":
    create_manual()
